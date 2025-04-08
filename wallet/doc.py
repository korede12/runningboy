# If you haven't installed python-docx, you may need to install it using:
# pip install python-docx

from docx import Document
from docx.shared import Pt

# Create a new document
document = Document()

# Define a style for normal text
style = document.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
document.add_heading('Lessons from Ramadan (دروس من رمضان)', 0)

# Section 1: Self-Control
document.add_heading('1. Self-Control (ضبط النفس)', level=1)

document.add_paragraph('📖 Qur’an:')
document.add_paragraph(
    'يَا أَيُّهَا الَّذِينَ آمَنُوا كُتِبَ عَلَيْكُمُ الصِّيَامُ كَمَا كُتِبَ عَلَى الَّذِينَ مِنْ قَبْلِكُمْ لَعَلَّكُمْ تَتَّقُونَ\n'
    'Yā ayyuhalladhīna āmanū kutiba ʿalaykumuṣ-ṣiyāmu kamā kutiba ʿala alladhīna min qablikum laʿallakum tattaqūn.\n'
    '"O you who have believed, fasting has been prescribed for you as it was prescribed for those before you, so that you may attain Taqwa (self-restraint)."'
    '\n(Surah Al-Baqarah 2:183)'
)

document.add_paragraph('🕋 Hadith:')
document.add_paragraph(
    'لَيْسَ الشَّدِيدُ بِالصُّرَعَةِ إِنَّمَا الشَّدِيدُ الَّذِي يَمْلِكُ نَفْسَهُ عِنْدَ الْغَضَبِ\n'
    'Laysa ash-shadīdu biṣ-ṣuraʿah, innama ash-shadīdu alladhī yamliku nafsahu ʿinda al-ghaḍab.\n'
    '"The strong is not the one who overcomes others by his strength, but the strong is the one who controls himself in anger."'
    '\n(Sahih al-Bukhari 6114, Sahih Muslim 2609)'
)

document.add_heading('Explanation:', level=2)
document.add_paragraph(
    "Self-control is one of the greatest lessons of Ramadan. The ability to govern one’s desires and impulses is a mark of true strength. "
    "Ramadan serves as a spiritual training ground, helping us discipline our nafs (inner self) by teaching us to delay gratification, resist temptation, and master our impulses."
)

document.add_paragraph('1. Resisting the Urge for Instant Gratification:')
document.add_paragraph(
    "In a world driven by instant pleasure, fasting trains us to pause and reflect. The hunger we feel is not just physical—it is a reminder of our ability to control ourselves. "
    "Just as a skilled horseman tames his horse with reins, fasting helps us rein in our desires and keep them under control.\n\n"
    "A man who lacks self-control is not truly free. True bondage is being enslaved to one’s desires—to eat whenever hunger strikes, to act upon every impulse, and to submit to temptations without resistance. "
    "Ramadan liberates us from this bondage, making us masters of ourselves rather than slaves to our cravings."
)

document.add_paragraph('2. Strengthening Willpower and Taming the Nafs:')
document.add_paragraph(
    "When we deny ourselves food and drink during the day, we prove that we have power over our physical urges. But self-control extends beyond hunger—it also applies to our emotions, speech, and actions.\n"
    "- We control our tongues from lying, gossiping, and backbiting.\n"
    "- We control our tempers, responding with patience instead of anger.\n"
    "- We control our eyes, lowering our gaze and avoiding what is impermissible.\n"
    "- We control our hearts, purifying them from jealousy, arrogance, and greed.\n\n"
    "This training teaches us that if we can resist the temptation of food and water—the most basic human needs—then we can resist anything that leads us away from righteousness."
)

document.add_paragraph('3. Overcoming Destructive Desires:')
document.add_paragraph(
    "Without self-control, a person is at the mercy of their lower self, easily led into destruction through unchecked desires. "
    "The Prophet ﷺ warned us: \"The worst vessel a human being can fill is his stomach.\" (Sunan Ibn Majah 3349)\n\n"
    "Many sins—greed, lust, anger—stem from a lack of self-discipline. Ramadan helps us conquer these weaknesses by strengthening our resolve, teaching us that true power is not in conquering others but in conquering oneself."
)

document.add_paragraph('4. The True Meaning of Freedom:')
document.add_paragraph(
    "Some believe that freedom means doing whatever one desires. In reality, a person controlled by their impulses is not free at all—they are trapped in a cycle of indulgence and regret.\n"
    "- A person addicted to food is a slave to hunger.\n"
    "- A person who cannot control their gaze is a prisoner of lust.\n"
    "- A person who cannot hold their tongue is a servant of anger.\n\n"
    "True freedom comes from self-mastery, from breaking free of desires that pull us away from Allah. The one who controls his desires controls his destiny."
)

document.add_paragraph(
    "Ramadan trains us to be stronger than our impulses, preparing us not just for fasting but for life itself. A person who learns self-discipline during Ramadan will find it easier to resist sin, delay gratification, and make wiser choices.\n\n"
    "Self-control is not just about refraining from food—it is about becoming a master of your own soul.\n\n"
    "May Allah grant us strength over our nafs and make us among those who control their desires for His sake. Ameen. 🤲✨"
)

# Section 2: Community Building and Unity
document.add_heading('2. Community Building and Unity (بناء المجتمع والوحدة)', level=1)

document.add_paragraph('📖 Qur’an:')
document.add_paragraph(
    'وَاعْتَصِمُوا بِحَبْلِ اللَّهِ جَمِيعًا وَلَا تَفَرَّقُوا\n'
    'Wa-ʿtaṣimū bi-ḥabli llāhi jamīʿan wa lā tafarraqū.\n'
    '"And hold firmly to the rope of Allah all together and do not become divided."'
    '\n(Surah Aal-E-Imran 3:103)'
)

document.add_paragraph('🕋 Hadith:')
document.add_paragraph(
    'لَيْسَ الْمُؤْمِنُ الَّذِي يَشْبَعُ وَجَارُهُ جَائِعٌ إِلَى جَنْبِهِ\n'
    'Laysa al-mu’minu alladhī yashbaʿu wa jāruhu jāiʿun ilā janbih.\n'
    '"The believer is not the one who eats his fill while his neighbor goes hungry."'
    '\n(Sunan al-Kubra 19049)'
)

document.add_heading('Explanation:', level=2)
document.add_paragraph(
    "Ramadan is not just about personal transformation—it is also a month of unity, community, and collective growth. The shared experience of fasting, praying, and giving charity strengthens the bonds between individuals, reminding us that we are one Ummah—a single body that thrives on mutual care and compassion."
)

document.add_paragraph('1. Strengthening Brotherhood and Sisterhood:')
document.add_paragraph(
    "When we fast together, break our fasts side by side, and stand shoulder to shoulder in Taraweeh prayers, we dissolve the barriers that separate us—be they of wealth, nationality, or status. The rich and the poor, the young and the old, all stand as equals before Allah. This powerful unity fosters love, humility, and brotherhood."
)

document.add_paragraph('2. Increasing Generosity and Care for Others:')
document.add_paragraph(
    "Ramadan softens the heart and amplifies our generosity. Through Sadaqah (charity) and Zakat (obligatory almsgiving), we are reminded of our duty to those in need. We begin to see the struggles of the less fortunate not as distant problems, but as urgent responsibilities.\n\n"
    "Fasting allows us to experience, even if temporarily, the hunger that millions endure daily. This firsthand experience increases our empathy and inspires us to give more, share more, and support one another without hesitation.\n\n"
    "The Prophet ﷺ said: \"The best of people are those who are most beneficial to others.\" (Musnad Ahmad 23408)"
)

document.add_paragraph('3. Removing Division and Discrimination:')
document.add_paragraph(
    "Ramadan is a time when differences in race, background, and social status become insignificant. In the masjid, the businessman prays next to the laborer, and the scholar stands beside the student. This month serves as a reminder that our worth is not in our wealth or worldly titles but in our piety and righteousness.\n\n"
    "Allah says: إِنَّ أَكْرَمَكُمْ عِندَ اللَّهِ أَتْقَاكُمْ (Inna akramakum ‘indallāhi atqākum) – \"Indeed, the most noble of you in the sight of Allah is the most righteous among you.\" (Surah Al-Hujurat 49:13)"
)

document.add_paragraph('4. Reviving the Spirit of Togetherness:')
document.add_paragraph(
    "Ramadan reignites the spirit of family, friendship, and social connection. Families come together for Suhoor and Iftar, neighbors share meals, and entire communities unite in acts of worship. The collective participation in fasting and prayer strengthens relationships and reminds us that we are never alone in our journey toward Allah.\n\n"
    "Even the simplest acts, like greeting one another with salaam or making dua for a fellow Muslim, contribute to this atmosphere of love and unity."
)

document.add_paragraph('5. The Month of Peace and Forgiveness:')
document.add_paragraph(
    "Ramadan is an opportunity to mend broken relationships and let go of past grudges. The Prophet ﷺ encouraged reconciliation, saying: \"It is not lawful for a Muslim to forsake his brother for more than three nights... and the best of them is the one who greets the other first.\" (Sahih al-Bukhari 6077)\n\n"
    "During this blessed month, we should seek forgiveness from those we have wronged and forgive those who have wronged us. By doing so, we cleanse our hearts and align ourselves with the spirit of Ramadan—mercy, kindness, and unity."
)

document.add_paragraph(
    "Ramadan transforms our hearts and communities. It teaches us that unity is not just about being together physically but about fostering love, generosity, and selflessness. When we embrace the values of Ramadan—compassion, generosity, and forgiveness—we strengthen not just ourselves but the entire Ummah.\n\n"
    "May this Ramadan bring peace, unity, and love to all hearts. Ameen. 🤲✨"
)

# Section 3: The Month of Grace
document.add_heading('3. The Month of Grace (شَهْرُ الرَّحْمَةِ وَالْفَضْلِ)', level=1)

document.add_paragraph('📖 Qur’an:')
document.add_paragraph(
    'شَهْرُ رَمَضَانَ الَّذِي أُنزِلَ فِيهِ الْقُرْآنُ هُدًى لِّلنَّاسِ وَبَيِّنَاتٍ مِّنَ الْهُدَىٰ وَالْفُرْقَانِ\n'
    'Shahru Ramaḍāna alladhī unzila fīhi al-Qurʾānu hudan lil-nāsi wa bayyinātin mina al-hudā wa al-furqān.\n'
    '"The month of Ramadan is the one in which the Qur’an was revealed, as a guidance for mankind and clear proofs of guidance and criterion (between right and wrong)."'
    '\n(Surah Al-Baqarah 2:185)'
)

document.add_paragraph('🕋 Hadith:')
document.add_paragraph(
    'إِذَا جَاءَ رَمَضَانُ فُتِحَتْ أَبْوَابُ الْجَنَّةِ، وَغُلِّقَتْ أَبْوَابُ النَّارِ، وَسُلْسِلَتِ الشَّيَاطِينُ\n'
    'Idhā jā’a Ramaḍānu futihat abwābu al-jannah, wa ghulliqat abwābu al-nār, wa sulsilati ash-shayāṭīn.\n'
    '"When Ramadan begins, the gates of Paradise are opened, the gates of Hell are closed, and the devils are chained."'
    '\n(Sahih al-Bukhari 1899, Sahih Muslim 1079)'
)

document.add_heading('Explanation:', level=2)
document.add_paragraph(
    "Ramadan is a month unlike any other—a time of divine mercy, abundant blessings, and spiritual renewal. It is a month where the doors of Jannah are flung open, the gates of Jahannam are sealed, and the shayateen are restrained, giving us an unparalleled opportunity to seek Allah’s forgiveness and elevate our souls.\n\n"
    "Allah, in His infinite mercy, has chosen Ramadan as a season of grace, where His kindness and blessings descend upon His servants. The rewards for good deeds are multiplied, sins are forgiven, and hearts are purified."
)

document.add_paragraph('1. The Blessings of Health and Strength:')
document.add_paragraph(
    "Fasting is not just a spiritual act; it is also a means of physical purification. Modern science has confirmed what Islam taught over 1400 years ago—fasting detoxifies the body, improves digestion, boosts immunity, and strengthens mental clarity.\n\n"
    "- Fasting removes toxins from the body.\n"
    "- It improves insulin sensitivity, reducing the risk of diabetes.\n"
    "- It strengthens self-discipline, making us more mindful of what we consume.\n\n"
    "The Prophet ﷺ said: \"Fast, and you shall be healthy.\" (Musnad Ahmad 8838)\n\n"
    "Beyond the body, Ramadan cleanses the heart and mind. A person who enters Ramadan with a burdened soul often leaves it with peace, clarity, and renewed purpose."
)

document.add_paragraph('2. Inner Peace and Mental Clarity:')
document.add_paragraph(
    "In our fast-paced world, where stress and anxiety have become the norm, Ramadan provides an oasis of calm. By stepping away from distractions and focusing on worship, we find a sense of inner peace that no worldly pleasure can provide.\n\n"
    "- The recitation of the Qur’an soothes the soul.\n"
    "- The nightly Taraweeh prayers bring tranquility to the heart.\n"
    "- Increased dua and dhikr fill us with hope and serenity.\n\n"
    "Allah promises: أَلَا بِذِكْرِ اللَّهِ تَطْمَئِنُّ الْقُلُوبُ (Ala bidhikri llāhi tatma’innu al-qulūb) – \"Verily, in the remembrance of Allah do hearts find rest.\" (Surah Ar-Ra’d 13:28)"
)

document.add_paragraph('3. Unlimited Mercy and Forgiveness:')
document.add_paragraph(
    "One of the greatest blessings of Ramadan is the chance for complete forgiveness. The Prophet ﷺ said: \"Whoever fasts during Ramadan with faith and seeking reward, his past sins will be forgiven.\" (Sahih al-Bukhari 38, Sahih Muslim 760)\n\n"
    "In this month, even the smallest act of goodness is rewarded immensely. A single prayer, a moment of repentance, or a sincere tear shed in dua could erase years of sins.\n\n"
    "The Prophet ﷺ also said: \"Every night of Ramadan, Allah frees people from Hellfire.\" (Sunan at-Tirmidhi 682)\n\n"
    "Every moment of Ramadan holds the potential for eternal salvation."
)

document.add_paragraph('4. The Qur’an: The Ultimate Gift of Ramadan:')
document.add_paragraph(
    "The greatest blessing of Ramadan is that it is the month in which the Qur’an was revealed—a guidance that illuminates the path to Jannah. During this month, believers immerse themselves in its recitation, reflecting on its verses and allowing its wisdom to transform their hearts.\n\n"
    "The Prophet ﷺ said: \"The best of you are those who learn the Qur’an and teach it to others.\" (Sahih al-Bukhari 5027)"
)

document.add_paragraph('5. The Night of Ultimate Grace: Laylat al-Qadr')
document.add_paragraph(
    "One of the most profound blessings of Ramadan is Laylat al-Qadr (The Night of Decree)—a single night better than a thousand months.\n\n"
    "Allah says: لَيْلَةُ الْقَدْرِ خَيْرٌ مِّنْ أَلْفِ شَهْرٍ (Laylatu al-Qadri khayrun min alfi shahr) – \"The Night of Decree is better than a thousand months.\" (Surah Al-Qadr 97:3)\n\n"
    "On this night, the destiny of souls is written, the angels descend, and duas are accepted. A single sincere prayer on Laylat al-Qadr could be equivalent to a lifetime of worship."
)

document.add_paragraph('Conclusion:')
document.add_paragraph(
    "Ramadan is not just another month—it is a divine opportunity to change our lives forever. It is a time of physical purification, mental renewal, and spiritual rebirth. It is a month where grace flows abundantly, sins are erased, and the gates of Jannah are open wide.\n\n"
    "To truly benefit from Ramadan, we must:\n"
    "✅ Seek forgiveness with sincerity.\n"
    "✅ Embrace the Qur’an as our guide.\n"
    "✅ Give generously and serve others.\n"
    "✅ Make the most of Laylat al-Qadr.\n\n"
    "May Allah make this Ramadan a means of elevating our faith, purifying our hearts, and granting us His infinite mercy. Ameen. 🤲✨"
)

# Section 4: Laylat al-Qadr – The Night of Majesty
document.add_heading('4. Laylat al-Qadr – The Night of Majesty (ليلة القدر)', level=1)

document.add_paragraph('📖 Qur’an:')
document.add_paragraph(
    'لَيْلَةُ الْقَدْرِ خَيْرٌ مِّنْ أَلْفِ شَهْرٍ\n'
    'Laylatu al-Qadri khayrun min alfi shahr.\n'
    '"The Night of Decree is better than a thousand months."'
    '\n(Surah Al-Qadr 97:3)'
)

document.add_paragraph('🕋 Hadith:')
document.add_paragraph(
    'مَنْ قَامَ لَيْلَةَ الْقَدْرِ إِيمَانًا وَاحْتِسَابًا غُفِرَ لَهُ مَا تَقَدَّمَ مِنْ ذَنْبِهِ\n'
    'Man qāma laylata l-qadri īmānan waḥtisāban ghufira lahu mā taqaddama min dhanbih.\n'
    '"Whoever stands in prayer during Laylat al-Qadr out of faith and in hope of reward, his previous sins will be forgiven."'
    '\n(Sahih al-Bukhari 1901, Sahih Muslim 760)'
)

document.add_heading('Explanation:', level=2)
document.add_paragraph(
    "Among the greatest blessings of Ramadan is Laylat al-Qadr—the Night of Majesty—found within the last ten days of the month. It is the night when destinies are written, sins are forgiven, and divine mercy is granted. "
    "Worship on this single night is more valuable than a thousand months. It is a time of unparalleled spiritual transformation, making it a golden opportunity for sincere supplication and devotion."
)

document.add_paragraph('Conclusion:')
document.add_paragraph(
    "Ramadan is a sacred period of immense spiritual growth, self-discipline, and communal harmony. It teaches us to master our desires, strengthens our bonds with others, and offers a night that changes destinies.\n\n"
    "May Allah bless us with the wisdom to embrace His guidance and the strength to carry His light beyond this blessed month. Ameen."
)

# Save the document
document.save('Ramadan_Lessons.docx')
print("The file 'Ramadan_Lessons.docx' has been created successfully.")
